"""
Block Extractor - Structured Content Extraction
================================================

Extracts blocks from manuscript files with inline styling marks.

Supports:
- 14 block types (front matter, chapters, paragraphs, scene breaks, back matter)
- 4 inline marks (italic, bold, smallcaps, code)
- DOCX, PDF, TXT formats

Author: Pronto Publishing
Version: 4.2.0 (Contract v1.1)

Contract v1.1 changes (2026-04-23):
  - Always emit `spans` for text-carrying blocks; never emit bare `text`.
    Plain text becomes a single span with empty marks.
  - Add `list` block detection with `meta.list_type` (ordered|unordered) and
    `meta.list_group` (an integer that groups consecutive list blocks so
    Worker 2 can wrap them in a single itemize/enumerate).
  - Ensure `meta.chapter_number` is always present on chapter_heading blocks:
    int when the heading text contains a number, `None` for unnumbered
    chapters (fixes a latent clobber bug where the caller was overwriting
    the detected number with a simple counter).
  - Fix REVIEW_NOTES C1: the generic front-matter and back-matter fallbacks
    no longer misclassify every unmatched paragraph as `front_matter_dedication`
    or `back_matter_about_author`. Unmatched body text now returns `paragraph`.

Out of scope for this change (tracked separately):
  - Prologue/Epilogue pattern detection (current CHAPTER_PATTERNS require a
    digit; unnumbered chapter titles still flow through as paragraphs).
  - Sub-section `heading` blocks with `meta.level` (current manuscripts don't
    use these; the downstream output_validator will surface it if they appear).
"""

import re
import logging
from typing import List, Dict, Any, Tuple, Optional
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
import PyPDF2

logger = logging.getLogger(__name__)


class BlockExtractor:
    """Extracts structured blocks from manuscript files."""
    
    # Chapter heading patterns
    CHAPTER_PATTERNS = [
        r'^Chapter\s+\d+',
        r'^CHAPTER\s+\d+',
        r'^Ch\.\s+\d+',
        r'^\d+\.',  # "1.", "2.", etc.
        r'^Part\s+\d+',
        r'^PART\s+\d+',
    ]
    
    # Front matter keywords
    FRONT_MATTER_KEYWORDS = {
        'dedication': 'front_matter_dedication',
        'copyright': 'front_matter_copyright',
        'title': 'front_matter_title',
        'contents': 'toc_marker',
        'table of contents': 'toc_marker',
    }
    
    # Back matter keywords
    BACK_MATTER_KEYWORDS = {
        'about the author': 'back_matter_about_author',
        'about author': 'back_matter_about_author',
        'also by': 'back_matter_also_by',
    }
    
    # Scene break patterns
    SCENE_BREAK_PATTERNS = [
        r'^\s*\*\s*\*\s*\*\s*$',  # * * *
        r'^\s*#\s*$',              # #
        r'^\s*~\s*$',              # ~
        r'^\s*-{3,}\s*$',          # ---
    ]

    # List-item patterns. Applied only if a line did NOT match a chapter
    # pattern first, so "1. Chapter Title" still classifies as chapter_heading.
    # A "1) Step one" style line will classify as a list.
    LIST_PATTERNS = [
        (r'^\s*[\u2022\u2023\u25E6\u2043\u2219]\s+', 'unordered'),  # bullet chars
        (r'^\s*[-*+]\s+',                              'unordered'),  # markdown-style
        (r'^\s*\d+[)]\s+',                             'ordered'),    # "1) foo"
        (r'^\s*[a-z][)]\s+',                           'ordered'),    # "a) foo"
    ]

    def __init__(self):
        """Initialize block extractor."""
        self.block_counter = 0
    
    def extract(self, file_path: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """
        Extract blocks from manuscript file.
        
        Args:
            file_path: Path to manuscript file
            
        Returns:
            Tuple of (blocks, source_meta)
        """
        # Reset block counter for each file
        self.block_counter = 0
        
        ext = Path(file_path).suffix.lower()
        
        if ext == '.docx':
            return self._extract_docx(file_path)
        elif ext == '.pdf':
            return self._extract_pdf(file_path)
        elif ext in ['.txt', '.md']:
            return self._extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _generate_block_id(self) -> str:
        """Generate unique block ID."""
        self.block_counter += 1
        return f"b_{self.block_counter:06d}"
    
    def _extract_docx(self, file_path: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Extract blocks from DOCX file with inline styling."""
        doc = Document(file_path)
        blocks = []

        # Track state
        in_front_matter = True
        in_back_matter = False
        chapter_count = 0
        list_group_counter = 0
        in_list = False

        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()

            if not text:
                continue  # Skip empty paragraphs

            # Detect block type
            block_type, meta = self._detect_block_type(
                text=text,
                style=para.style.name if para.style else None,
                in_front_matter=in_front_matter,
                in_back_matter=in_back_matter
            )

            # Update state
            if block_type == 'chapter_heading':
                chapter_count += 1
                in_front_matter = False
                # meta.chapter_number is authoritative when the detector
                # extracted one from the heading text; otherwise the chapter
                # is unnumbered (e.g. a future Prologue path) and we set None.
                if 'chapter_number' not in meta:
                    meta['chapter_number'] = None

            if block_type.startswith('front_matter') or block_type == 'toc_marker':
                in_front_matter = True

            if block_type.startswith('back_matter'):
                in_back_matter = True
                in_front_matter = False

            # list_group state — consecutive `list` blocks share a group.
            if block_type == 'list':
                if not in_list:
                    list_group_counter += 1
                    in_list = True
                meta['list_group'] = list_group_counter
            else:
                in_list = False

            # Extract spans with inline marks
            spans = self._extract_spans_from_para(para)

            # Build block. Contract v1.1: always spans, never bare `text`.
            # Plain text is a single span with empty marks.
            block = {
                'id': self._generate_block_id(),
                'type': block_type,
                'spans': spans if spans else [{'text': text, 'marks': []}],
            }

            # Add metadata if present
            if meta:
                block['meta'] = meta

            # Add source location
            block['source_loc'] = {
                'doc_paragraph_index': para_idx
            }

            blocks.append(block)
        
        # Source metadata
        source_meta = {
            'original_filename': Path(file_path).name,
            'original_format': 'docx',
            'total_paragraphs': len(doc.paragraphs),
            'detected_chapters': chapter_count,
            'has_front_matter': any(b['type'].startswith('front_matter') or b['type'] == 'toc_marker' for b in blocks),
            'has_back_matter': any(b['type'].startswith('back_matter') for b in blocks)
        }
        
        logger.info(f"Extracted {len(blocks)} blocks from DOCX ({chapter_count} chapters)")
        
        return blocks, source_meta
    
    def _extract_pdf(self, file_path: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Extract blocks from PDF file (text-only, no styling)."""
        blocks = []

        # Track state
        in_front_matter = True
        in_back_matter = False
        chapter_count = 0
        list_group_counter = 0
        in_list = False

        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            for page_num, page in enumerate(reader.pages, start=1):
                text = page.extract_text()
                lines = text.split('\n')

                for line in lines:
                    line = line.strip()

                    if not line:
                        continue

                    # Detect block type
                    block_type, meta = self._detect_block_type(
                        text=line,
                        style=None,
                        in_front_matter=in_front_matter,
                        in_back_matter=in_back_matter
                    )

                    # Update state
                    if block_type == 'chapter_heading':
                        chapter_count += 1
                        in_front_matter = False
                        if 'chapter_number' not in meta:
                            meta['chapter_number'] = None

                    if block_type.startswith('front_matter') or block_type == 'toc_marker':
                        in_front_matter = True

                    if block_type.startswith('back_matter'):
                        in_back_matter = True
                        in_front_matter = False

                    if block_type == 'list':
                        if not in_list:
                            list_group_counter += 1
                            in_list = True
                        meta['list_group'] = list_group_counter
                    else:
                        in_list = False

                    # Contract v1.1: always spans. PDFs have no inline marks.
                    block = {
                        'id': self._generate_block_id(),
                        'type': block_type,
                        'spans': [{'text': line, 'marks': []}],
                    }

                    if meta:
                        block['meta'] = meta

                    block['source_loc'] = {
                        'doc_page_number': page_num
                    }

                    blocks.append(block)
        
        # Source metadata
        source_meta = {
            'original_filename': Path(file_path).name,
            'original_format': 'pdf',
            'total_pages': len(reader.pages),
            'detected_chapters': chapter_count,
            'has_front_matter': any(b['type'].startswith('front_matter') or b['type'] == 'toc_marker' for b in blocks),
            'has_back_matter': any(b['type'].startswith('back_matter') for b in blocks)
        }
        
        logger.info(f"Extracted {len(blocks)} blocks from PDF ({chapter_count} chapters)")
        
        return blocks, source_meta
    
    def _extract_txt(self, file_path: str) -> Tuple[List[Dict[str, Any]], Dict[str, Any]]:
        """Extract blocks from plain text file."""
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        blocks = []

        # Track state
        in_front_matter = True
        in_back_matter = False
        chapter_count = 0
        list_group_counter = 0
        in_list = False

        for line_num, line in enumerate(lines, start=1):
            line = line.strip()

            if not line:
                continue

            # Detect block type
            block_type, meta = self._detect_block_type(
                text=line,
                style=None,
                in_front_matter=in_front_matter,
                in_back_matter=in_back_matter
            )

            # Update state
            if block_type == 'chapter_heading':
                chapter_count += 1
                in_front_matter = False
                if 'chapter_number' not in meta:
                    meta['chapter_number'] = None

            if block_type.startswith('front_matter') or block_type == 'toc_marker':
                in_front_matter = True

            if block_type.startswith('back_matter'):
                in_back_matter = True
                in_front_matter = False

            if block_type == 'list':
                if not in_list:
                    list_group_counter += 1
                    in_list = True
                meta['list_group'] = list_group_counter
            else:
                in_list = False

            # Contract v1.1: always spans. TXT has no inline marks.
            block = {
                'id': self._generate_block_id(),
                'type': block_type,
                'spans': [{'text': line, 'marks': []}],
            }

            if meta:
                block['meta'] = meta

            blocks.append(block)
        
        # Source metadata
        source_meta = {
            'original_filename': Path(file_path).name,
            'original_format': 'txt',
            'total_lines': len(lines),
            'detected_chapters': chapter_count,
            'has_front_matter': any(b['type'].startswith('front_matter') or b['type'] == 'toc_marker' for b in blocks),
            'has_back_matter': any(b['type'].startswith('back_matter') for b in blocks)
        }
        
        logger.info(f"Extracted {len(blocks)} blocks from TXT ({chapter_count} chapters)")
        
        return blocks, source_meta
    
    def _detect_block_type(
        self,
        text: str,
        style: Optional[str],
        in_front_matter: bool,
        in_back_matter: bool
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Detect block type from text and context.

        Returns:
            Tuple of (block_type, meta_dict). For `list` blocks, meta carries
            only `list_type`; the caller is responsible for assigning
            `list_group` since that requires cross-block state. For
            `chapter_heading` blocks with a numeric heading, meta carries
            `chapter_number`; otherwise the caller sets it to None.
        """
        text_lower = text.lower()
        meta: Dict[str, Any] = {}

        # Check for scene break
        for pattern in self.SCENE_BREAK_PATTERNS:
            if re.match(pattern, text):
                return 'scene_break', {'pattern': text}

        # Check for chapter heading (checked before list detection so a line
        # like "1. Chapter Title" classifies as chapter, not list).
        for pattern in self.CHAPTER_PATTERNS:
            if re.match(pattern, text, re.IGNORECASE):
                match = re.search(r'\d+', text)
                if match:
                    meta['chapter_number'] = int(match.group())
                return 'chapter_heading', meta

        # Check for list-item patterns (bullets, ordered markers).
        for pattern, list_type in self.LIST_PATTERNS:
            if re.match(pattern, text):
                return 'list', {'list_type': list_type}

        # Check for specific front matter types
        for keyword, block_type in self.FRONT_MATTER_KEYWORDS.items():
            if keyword in text_lower:
                return block_type, {'detected_keyword': keyword}

        # Check for specific back matter types
        for keyword, block_type in self.BACK_MATTER_KEYWORDS.items():
            if keyword in text_lower:
                return block_type, {'detected_keyword': keyword}

        # No specific match. Fall through to paragraph regardless of
        # in_front_matter / in_back_matter state. Previously this returned
        # front_matter_dedication / back_matter_about_author as a catch-all,
        # which produced the REVIEW_NOTES C1 bug (every paragraph before
        # Chapter 1 tagged as dedication). The canon's "no fabrication" rule
        # applies: if it doesn't match an explicit signal, it's a paragraph.
        return 'paragraph', {}
    
    def _extract_spans_from_para(self, para: Paragraph) -> List[Dict[str, Any]]:
        """
        Extract spans with inline marks from DOCX paragraph.
        
        Returns list of span objects: [{text: str, marks: [str]}]
        """
        spans = []
        
        for run in para.runs:
            run_text = run.text
            
            if not run_text:
                continue
            
            # Collect marks for this span
            marks = []
            
            if run.italic:
                marks.append('italic')
            
            if run.bold:
                marks.append('bold')
            
            if run.font and run.font.small_caps:
                marks.append('smallcaps')
            
            # Code detection (monospace font)
            if run.font and run.font.name and 'mono' in run.font.name.lower():
                marks.append('code')
            
            # Create span
            spans.append({
                'text': run_text,
                'marks': marks
            })
        
        return spans
